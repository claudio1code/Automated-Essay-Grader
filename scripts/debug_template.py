import os
import sys
import argparse
import re
from docx import Document

def analisar_documento(file_path: str):
    """
    Este script analisa um arquivo .docx para diagnosticar a estrutura do
    conteúdo, procurando por placeholders não substituídos e por artefatos
    de substituição incorreta (como chaves '{}' sozinhas).
    """
    try:
        document = Document(file_path)
        print(f"✅ Analisando o arquivo: {file_path}\n")
    except Exception as e:
        print(f"❌ Erro ao abrir o arquivo: {e}")
        return

    placeholders_regex = re.compile(r"\{\{.*?\}}")
    error_artifact_regex = re.compile(r"\{[^}]*?\}|^\s*\}\s*$") # Pega '{texto}' ou uma '}' sozinha na linha

    all_paragraphs = []
    for p in document.paragraphs:
        all_paragraphs.append(p)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)
    for section in document.sections:
        for p in section.header.paragraphs:
            all_paragraphs.append(p)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paragraphs.extend(cell.paragraphs)

    print("--- INICIANDO ANÁLISE DE CONTEÚDO ---")
    print("Procurando por placeholders restantes e por artefatos de erro (ex: '{40}' ou '}').\n")

    placeholders_found = []
    errors_found = []

    for i, p in enumerate(all_paragraphs):
        if not p.text.strip():
            continue

        found_placeholders = placeholders_regex.findall(p.text)
        if found_placeholders:
            placeholders_found.extend(found_placeholders)
            print(f"⚠️ Placeholder não substituído encontrado no parágrafo {i}: {found_placeholders}")
            print(f"   Texto: '{p.text.strip()}'\n")

        found_errors = error_artifact_regex.findall(p.text)
        # Filtra para não pegar os placeholders que já foram reportados
        filtered_errors = [err for err in found_errors if not err.startswith("{{ ")]
        if filtered_errors:
            errors_found.extend(filtered_errors)
            print(f"❌ ARTEFATO DE ERRO encontrado no parágrafo {i}: {filtered_errors}")
            print(f"   Texto: '{p.text.strip()}'")
            print(f"   Estrutura de Runs:")
            for r_idx, run in enumerate(p.runs):
                print(f"     - Run {r_idx}: '{run.text}'")
            print("")


    print("\n--- RESULTADO DA ANÁLISE ---")
    if not placeholders_found and not errors_found:
        print("✅ SUCESSO: Nenhum placeholder restante ou artefato de erro foi encontrado.")
        print("   Isso indica que todas as substituições provavelmente foram bem-sucedidas.")
    else:
        if placeholders_found:
            print(f"🚨 ENCONTRADOS {len(placeholders_found)} PLACEHOLDERS NÃO SUBSTITUÍDOS.")
        if errors_found:
            print(f"🚨 ENCONTRADOS {len(errors_found)} ARTEFATOS DE ERRO (como '{{}}').")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Analisa a estrutura e conteúdo de um arquivo .docx.")
    parser.add_argument("file_path", type=str, help="Caminho para o arquivo .docx a ser analisado.")
    args = parser.parse_args()
    
    analisar_documento(args.file_path)
