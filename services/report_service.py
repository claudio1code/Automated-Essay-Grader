from io import BytesIO
from typing import Any, Dict, Optional

from docx import Document
from docx.text.paragraph import Paragraph

from config import Config
from logger import get_logger

logger = get_logger(__name__)


def preencher_e_gerar_docx(
    dados_redacao: Dict[str, Any], caminho_template: str = Config.TEMPLATE_DOCX_PATH
) -> Optional[BytesIO]:
    """
    Preenche o template .docx preservando formatação e retorna o arquivo em bytes.

    Args:
        dados_redacao (Dict[str, Any]): Dicionário com os dados da correção.
        caminho_template (str): Caminho para o arquivo de template .docx.

    Returns:
        Optional[BytesIO]: Buffer contendo o documento gerado ou None em caso de erro.
    """
    try:
        logger.info(f"Gerando DOCX usando template: {caminho_template}")
        document = Document(caminho_template)

        # Acessa os dados aninhados de forma segura
        analise_comps = dados_redacao.get("analise_competencias", {})
        c1 = analise_comps.get("c1", {})
        c2 = analise_comps.get("c2", {})
        c3 = analise_comps.get("c3", {})
        c4 = analise_comps.get("c4", {})
        c5 = analise_comps.get("c5", {})

        # Dicionário de substituições
        substituicoes = {
            "{{NOME_ALUNO}}": dados_redacao.get("nome_aluno", "Não informado"),
            "{{TEMA}}": dados_redacao.get("tema_redacao", "Não informado"),
            "{{DATA}}": dados_redacao.get("data_redacao", "Não informada"),
            "{{NOTA_TOTAL}}": str(dados_redacao.get("nota_final", "N/A")),
            "{{COMENTARIOS}}": dados_redacao.get("comentarios_gerais", ""),
            "{{NOTA_C1}}": str(c1.get("nota", "N/A")),
            "{{ANALISE_C1}}": c1.get("analise", ""),
            "{{NOTA_C2}}": str(c2.get("nota", "N/A")),
            "{{ANALISE_C2}}": c2.get("analise", ""),
            "{{NOTA_C3}}": str(c3.get("nota", "N/A")),
            "{{ANALISE_C3}}": c3.get("analise", ""),
            "{{NOTA_C4}}": str(c4.get("nota", "N/A")),
            "{{ANALISE_C4}}": c4.get("analise", ""),
            "{{NOTA_C5}}": str(c5.get("nota", "N/A")),
            "{{ANALISE_C5}}": c5.get("analise", ""),
        }

        # Lógica para o alerta de originalidade
        alerta = dados_redacao.get("alerta_originalidade")
        if alerta:
            substituicoes["{{ALERTA_ORIGINALIDADE}}"] = (
                f"⚠️ ALERTA DE ORIGINALIDADE: {alerta}"
            )
        else:
            substituicoes["{{ALERTA_ORIGINALIDADE}}"] = ""

        # Substituir placeholders em parágrafos PRESERVANDO formatação
        for paragrafo in document.paragraphs:
            substituir_em_paragrafo(paragrafo, substituicoes)

        # Substituir placeholders em tabelas PRESERVANDO formatação
        for tabela in document.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in celula.paragraphs:
                        substituir_em_paragrafo(paragrafo, substituicoes)

        # Salva o documento em um buffer de memória
        doc_buffer = BytesIO()
        document.save(doc_buffer)
        doc_buffer.seek(0)

        logger.info("Arquivo DOCX gerado com sucesso.")
        return doc_buffer

    except FileNotFoundError:
        logger.critical(
            f"O arquivo de template '{caminho_template}' não foi encontrado."
        )
        return None
    except Exception as e:
        logger.error(f"Erro ao gerar o arquivo DOCX: {e}")
        return None


def substituir_em_paragrafo(
    paragrafo: Paragraph, substituicoes: Dict[str, str]
) -> None:
    """
    Substitui placeholders em um parágrafo preservando a formatação.

    Args:
        paragrafo (Paragraph): Objeto parágrafo do python-docx.
        substituicoes (Dict[str, str]): Dicionário de chaves e valores para substituição.
    """
    # Verifica se há algum placeholder no texto completo do parágrafo
    texto_completo = paragrafo.text

    # Otimização: se não tem chaves de template, retorna cedo
    if "{{" not in texto_completo:
        return

    for placeholder, valor in substituicoes.items():
        if placeholder in texto_completo:
            # Substitui preservando formatação dentro dos 'runs'
            for run in paragrafo.runs:
                if placeholder in run.text:
                    # str(valor) garante que números sejam convertidos para string
                    run.text = run.text.replace(placeholder, str(valor))
